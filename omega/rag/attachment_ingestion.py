"""Attachment ingestion helpers for PDF/DOCX/HTML/images with bounded chunking."""

from __future__ import annotations

from dataclasses import dataclass, field
import base64
import hashlib
from pathlib import Path
import io
import re
import tempfile
import threading
import json
import os
import signal
import subprocess
import sys
import zipfile
from typing import Any, Dict, List, Mapping, Optional, Sequence, Tuple

from omega.vision.contracts import (
    OCRQualityPolicy,
    OCRQualitySummary,
    OCRSpan,
    normalize_ocr_spans,
)
from omega.vision.ocr_runtime import OCRWorkerOverloadedError


DEFAULT_MAX_FILE_BYTES = 20 * 1024 * 1024
DEFAULT_MAX_EXTRACTED_CHARS = 200_000
DEFAULT_MAX_CHUNK_CHARS = 2_000
DEFAULT_CHUNK_OVERLAP = 200
DEFAULT_HIDDEN_PREFIX = "[hidden_html] "
DEFAULT_OCR_MAX_SPANS_PER_CHUNK = 8

MIME_TO_FORMAT = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/html": "html",
    "image/png": "image",
    "image/jpeg": "image",
    "image/jpg": "image",
    "image/webp": "image",
    "image/gif": "image",
}

HIDDEN_STYLE_RE = re.compile(
    r"(display\s*:\s*none|visibility\s*:\s*hidden)", flags=re.IGNORECASE
)
_PADDLE_OCR_CACHE: Dict[tuple[str, bool], Any] = {}
_PADDLE_OCR_CACHE_LOCK = threading.Lock()
_RAPIDOCR_CACHE: Dict[tuple[str], Any] = {}
_RAPIDOCR_CACHE_LOCK = threading.Lock()


@dataclass(frozen=True)
class AttachmentVisualAsset:
    """Parser-bound visual payload; consumed before ContentItem construction."""

    asset_id: str
    mime: str
    sha256: str
    payload_b64: str
    role: str
    source_kind: str
    page_number: Optional[int] = None
    embedded_index: Optional[int] = None
    width: int = 0
    height: int = 0
    size_bytes: int = 0

    def decode(self) -> bytes:
        raw = base64.b64decode(str(self.payload_b64), validate=True)
        if hashlib.sha256(raw).hexdigest() != str(self.sha256):
            raise ValueError("visual asset sha256 mismatch")
        return raw


@dataclass(frozen=True)
class AttachmentVisualConfig:
    enabled: bool = False
    render_pdf_pages: bool = True
    extract_embedded_images: bool = True
    max_assets: int = 8
    max_total_bytes: int = 32 * 1024 * 1024
    max_asset_bytes: int = 8 * 1024 * 1024
    max_pdf_pages: int = 8
    pdf_dpi: int = 120
    max_pdf_pixels_per_page: int = 16_000_000
    min_width: int = 32
    min_height: int = 32
    failure_policy: str = "degrade"


@dataclass(frozen=True)
class AttachmentChunk:
    text: str
    kind: str
    is_hidden: bool = False
    ocr_span_ids: List[str] = field(default_factory=list)
    char_start: Optional[int] = None
    char_end: Optional[int] = None


@dataclass(frozen=True)
class AttachmentExtractResult:
    text: str
    chunks: List[AttachmentChunk]
    format: str
    text_empty: bool
    scan_like: bool
    hidden_text_chars: int
    warnings: List[str]
    recommended_verdict: str
    is_image: bool = False
    image_mime: Optional[str] = None
    image_sha256: Optional[str] = None
    image_bytes_size: int = 0
    ocr_status: str = "none"
    ocr_provider: Optional[str] = None
    ocr_text_chars: int = 0
    ocr_spans: List[OCRSpan] = field(default_factory=list)
    ocr_quality: OCRQualitySummary = field(default_factory=OCRQualitySummary)
    visual_assets: List[AttachmentVisualAsset] = field(default_factory=list)
    visual_status: str = "none"


@dataclass(frozen=True)
class AttachmentOCRConfig:
    enabled: str = "false"
    provider: str = "rapidocr"
    execution_mode: str = "inline"
    prewarm: bool = False
    worker_startup_timeout_sec: float = 25.0
    worker_request_timeout_sec: float = 15.0
    worker_max_memory_mb: int = 2048
    worker_max_requests: int = 500
    worker_pool_size: int = 1
    worker_max_pending_requests: int = 2
    worker_queue_timeout_sec: float = 1.0
    worker_intra_op_threads: int = 2
    worker_inter_op_threads: int = 1
    lang: str = "en"
    use_angle_cls: bool = True
    max_text_chars: int = DEFAULT_MAX_EXTRACTED_CHARS
    max_spans_per_chunk: int = DEFAULT_OCR_MAX_SPANS_PER_CHUNK
    min_confidence: float = 0.50
    max_spans: int = 256
    max_span_chars: int = 512
    require_geometry: bool = False
    min_polygon_area_px: float = 0.0
    failure_policy: str = "degrade"


@dataclass(frozen=True)
class ImageOCRResult:
    text: str
    status: str
    provider: Optional[str]
    warnings: List[str] = field(default_factory=list)
    spans: List[OCRSpan] = field(default_factory=list)


@dataclass(frozen=True)
class AttachmentSandboxConfig:
    enabled: bool = False
    timeout_sec: int = 8
    max_memory_mb: int = 512
    max_cpu_sec: int = 6


@dataclass(frozen=True)
class AttachmentIngestionConfig:
    enabled: bool = True
    max_file_bytes: int = DEFAULT_MAX_FILE_BYTES
    max_extracted_chars: int = DEFAULT_MAX_EXTRACTED_CHARS
    max_chunk_chars: int = DEFAULT_MAX_CHUNK_CHARS
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP
    html_include_hidden: bool = True
    hidden_chunk_prefix: str = DEFAULT_HIDDEN_PREFIX
    scan_like_min_chars_per_page: int = 25
    scan_like_min_alpha_ratio: float = 0.30
    zip_enabled: bool = False
    zip_max_files: int = 100
    zip_max_depth: int = 5
    zip_max_total_bytes: int = 20 * 1024 * 1024
    zip_allow_encrypted: bool = False
    strict_magic: bool = False
    max_pdf_pages: int = 250
    max_docx_entries: int = 2000
    max_docx_uncompressed_bytes: int = 64 * 1024 * 1024
    max_image_pixels: int = 40_000_000
    max_html_nodes: int = 200_000
    sandbox: AttachmentSandboxConfig = field(default_factory=AttachmentSandboxConfig)
    ocr: AttachmentOCRConfig = field(default_factory=AttachmentOCRConfig)
    visual: AttachmentVisualConfig = field(default_factory=AttachmentVisualConfig)

    @classmethod
    def from_cfg(cls, cfg: Mapping[str, Any] | None) -> "AttachmentIngestionConfig":
        data = dict(cfg or {})
        zip_cfg = data.get("zip", {}) if isinstance(data.get("zip", {}), dict) else {}
        ocr_cfg = data.get("ocr", {}) if isinstance(data.get("ocr", {}), dict) else {}
        sandbox_cfg = (
            data.get("sandbox", {}) if isinstance(data.get("sandbox", {}), dict) else {}
        )
        visual_cfg = (
            data.get("visual", {}) if isinstance(data.get("visual", {}), dict) else {}
        )
        parser_limits = (
            data.get("parser_limits", {})
            if isinstance(data.get("parser_limits", {}), dict)
            else {}
        )
        max_chunk = int(data.get("max_chunk_chars", DEFAULT_MAX_CHUNK_CHARS))
        overlap = int(data.get("chunk_overlap", DEFAULT_CHUNK_OVERLAP))
        if max_chunk <= 0:
            max_chunk = DEFAULT_MAX_CHUNK_CHARS
        if overlap < 0:
            overlap = DEFAULT_CHUNK_OVERLAP
        if overlap >= max_chunk:
            overlap = max_chunk - 1
        return cls(
            enabled=bool(data.get("enabled", True)),
            max_file_bytes=int(data.get("max_file_bytes", DEFAULT_MAX_FILE_BYTES)),
            max_extracted_chars=int(
                data.get("max_extracted_chars", DEFAULT_MAX_EXTRACTED_CHARS)
            ),
            max_chunk_chars=max_chunk,
            chunk_overlap=overlap,
            html_include_hidden=bool(data.get("html_include_hidden", True)),
            hidden_chunk_prefix=str(
                data.get("hidden_chunk_prefix", DEFAULT_HIDDEN_PREFIX)
            ),
            scan_like_min_chars_per_page=int(
                data.get("scan_like_min_chars_per_page", 25)
            ),
            scan_like_min_alpha_ratio=float(
                data.get("scan_like_min_alpha_ratio", 0.30)
            ),
            zip_enabled=bool(zip_cfg.get("enabled", False)),
            zip_max_files=int(zip_cfg.get("max_files", 100)),
            zip_max_depth=int(zip_cfg.get("max_depth", 5)),
            zip_max_total_bytes=int(zip_cfg.get("max_total_bytes", 20 * 1024 * 1024)),
            zip_allow_encrypted=bool(zip_cfg.get("allow_encrypted", False)),
            strict_magic=bool(data.get("strict_magic", False)),
            max_pdf_pages=max(1, int(parser_limits.get("max_pdf_pages", 250))),
            max_docx_entries=max(1, int(parser_limits.get("max_docx_entries", 2000))),
            max_docx_uncompressed_bytes=max(
                1,
                int(parser_limits.get("max_docx_uncompressed_bytes", 64 * 1024 * 1024)),
            ),
            max_image_pixels=max(
                1, int(parser_limits.get("max_image_pixels", 40_000_000))
            ),
            max_html_nodes=max(1, int(parser_limits.get("max_html_nodes", 200_000))),
            sandbox=AttachmentSandboxConfig(
                enabled=bool(sandbox_cfg.get("enabled", False)),
                timeout_sec=max(1, int(sandbox_cfg.get("timeout_sec", 8))),
                max_memory_mb=max(64, int(sandbox_cfg.get("max_memory_mb", 512))),
                max_cpu_sec=max(1, int(sandbox_cfg.get("max_cpu_sec", 6))),
            ),
            ocr=AttachmentOCRConfig(
                enabled=str(ocr_cfg.get("enabled", "false")).strip().lower() or "false",
                provider=str(ocr_cfg.get("provider", "rapidocr")).strip().lower()
                or "rapidocr",
                execution_mode=str(ocr_cfg.get("execution_mode", "inline"))
                .strip()
                .lower()
                or "inline",
                prewarm=bool(ocr_cfg.get("prewarm", False)),
                worker_startup_timeout_sec=float(
                    ocr_cfg.get("worker_startup_timeout_sec", 25.0)
                ),
                worker_request_timeout_sec=float(
                    ocr_cfg.get("worker_request_timeout_sec", 15.0)
                ),
                worker_max_memory_mb=max(
                    256, int(ocr_cfg.get("worker_max_memory_mb", 2048))
                ),
                worker_max_requests=max(
                    1, int(ocr_cfg.get("worker_max_requests", 500))
                ),
                worker_pool_size=max(1, int(ocr_cfg.get("worker_pool_size", 1))),
                worker_max_pending_requests=max(
                    0, int(ocr_cfg.get("worker_max_pending_requests", 2))
                ),
                worker_queue_timeout_sec=float(
                    ocr_cfg.get("worker_queue_timeout_sec", 1.0)
                ),
                worker_intra_op_threads=max(
                    1, min(16, int(ocr_cfg.get("worker_intra_op_threads", 2)))
                ),
                worker_inter_op_threads=max(
                    1, min(8, int(ocr_cfg.get("worker_inter_op_threads", 1)))
                ),
                lang=str(ocr_cfg.get("lang", "en")).strip() or "en",
                use_angle_cls=bool(ocr_cfg.get("use_angle_cls", True)),
                max_text_chars=int(
                    ocr_cfg.get("max_text_chars", DEFAULT_MAX_EXTRACTED_CHARS)
                ),
                max_spans_per_chunk=max(
                    1,
                    int(
                        ocr_cfg.get(
                            "max_spans_per_chunk", DEFAULT_OCR_MAX_SPANS_PER_CHUNK
                        )
                    ),
                ),
                min_confidence=float(ocr_cfg.get("min_confidence", 0.50)),
                max_spans=max(1, int(ocr_cfg.get("max_spans", 256))),
                max_span_chars=max(1, int(ocr_cfg.get("max_span_chars", 512))),
                require_geometry=bool(ocr_cfg.get("require_geometry", False)),
                min_polygon_area_px=float(ocr_cfg.get("min_polygon_area_px", 0.0)),
                failure_policy=str(ocr_cfg.get("failure_policy", "degrade"))
                .strip()
                .lower()
                or "degrade",
            ),
            visual=AttachmentVisualConfig(
                enabled=bool(visual_cfg.get("enabled", False)),
                render_pdf_pages=bool(visual_cfg.get("render_pdf_pages", True)),
                extract_embedded_images=bool(
                    visual_cfg.get("extract_embedded_images", True)
                ),
                max_assets=max(1, int(visual_cfg.get("max_assets", 8))),
                max_total_bytes=max(
                    1, int(visual_cfg.get("max_total_bytes", 32 * 1024 * 1024))
                ),
                max_asset_bytes=max(
                    1, int(visual_cfg.get("max_asset_bytes", 8 * 1024 * 1024))
                ),
                max_pdf_pages=max(1, int(visual_cfg.get("max_pdf_pages", 8))),
                pdf_dpi=max(72, min(300, int(visual_cfg.get("pdf_dpi", 120)))),
                max_pdf_pixels_per_page=max(
                    1, int(visual_cfg.get("max_pdf_pixels_per_page", 16_000_000))
                ),
                min_width=max(1, int(visual_cfg.get("min_width", 32))),
                min_height=max(1, int(visual_cfg.get("min_height", 32))),
                failure_policy=str(visual_cfg.get("failure_policy", "degrade"))
                .strip()
                .lower()
                or "degrade",
            ),
        )


def _detect_format(*, path: Path | None, filename: str | None, mime: str | None) -> str:
    if mime:
        mime_l = str(mime).strip().lower()
        if mime_l in MIME_TO_FORMAT:
            return MIME_TO_FORMAT[mime_l]
    candidates: List[str] = []
    if path is not None:
        candidates.append(path.suffix.lower())
    if filename:
        candidates.append(Path(filename).suffix.lower())
    ext = next((x for x in candidates if x), "")
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext in {".html", ".htm"}:
        return "html"
    if ext in {".png", ".jpg", ".jpeg", ".webp", ".gif"}:
        return "image"
    if ext == ".zip":
        return "zip"
    return "text"


def _load_raw_bytes(*, path: Path | None, content_bytes: bytes | None) -> bytes:
    if content_bytes is not None:
        return content_bytes
    if path is None:
        raise ValueError("either path or content_bytes must be provided")
    return path.read_bytes()


def _normalize_text(text: str) -> str:
    return " ".join(str(text).split()).strip()


def _sha256_hex(raw: bytes) -> str:
    import hashlib

    h = hashlib.sha256()
    h.update(bytes(raw))
    return h.hexdigest()


def _chunk_text_ranges(
    text: str, *, max_chunk_chars: int, chunk_overlap: int
) -> List[Tuple[str, int, int]]:
    text_norm = _normalize_text(text)
    if not text_norm:
        return []
    if len(text_norm) <= max_chunk_chars:
        return [(text_norm, 0, len(text_norm))]

    step = max(1, max_chunk_chars - chunk_overlap)
    chunks: List[Tuple[str, int, int]] = []
    i = 0
    while i < len(text_norm):
        j = min(len(text_norm), i + max_chunk_chars)
        chunk = text_norm[i:j].strip()
        if chunk:
            offset = text_norm[i:j].find(chunk)
            start = i + max(0, offset)
            end = start + len(chunk)
            chunks.append((chunk, start, end))
        if j >= len(text_norm):
            break
        i += step
    return chunks


def _chunk_text(text: str, *, max_chunk_chars: int, chunk_overlap: int) -> List[str]:
    return [
        chunk
        for chunk, _, _ in _chunk_text_ranges(
            text, max_chunk_chars=max_chunk_chars, chunk_overlap=chunk_overlap
        )
    ]


def _clip_text(text: str, *, max_chars: int, warnings: List[str]) -> str:
    if len(text) <= max_chars:
        return text
    warnings.append("max_extracted_chars_truncated")
    return text[:max_chars]


def _missing_dependency_error(feature: str, dependency: str) -> RuntimeError:
    return RuntimeError(
        f"{feature} ingestion requires optional dependency '{dependency}'. "
        "Install with: pip install -e .[attachments]"
    )


def _extract_pdf_text_from_bytes(raw: bytes) -> Tuple[str, int]:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - dependency gate
        raise _missing_dependency_error("PDF", "pypdf") from exc

    reader = PdfReader(io.BytesIO(raw))
    pages: List[str] = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages), len(reader.pages)


def _extract_docx_text_from_bytes(raw: bytes) -> str:
    try:
        import docx  # type: ignore
    except Exception as exc:  # pragma: no cover - dependency gate
        raise _missing_dependency_error("DOCX", "python-docx") from exc

    document = docx.Document(io.BytesIO(raw))
    parts: List[str] = []

    for p in document.paragraphs:
        t = _normalize_text(p.text)
        if t:
            parts.append(t)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                t = _normalize_text(cell.text)
                if t:
                    parts.append(t)

    for section in document.sections:
        for p in section.header.paragraphs:
            t = _normalize_text(p.text)
            if t:
                parts.append(t)
        for p in section.footer.paragraphs:
            t = _normalize_text(p.text)
            if t:
                parts.append(t)

    return "\n".join(parts)


def _element_is_hidden(tag: Any) -> bool:
    style = str(tag.attrs.get("style", "") or "")
    if HIDDEN_STYLE_RE.search(style):
        return True
    if tag.has_attr("hidden"):
        return True
    if str(tag.attrs.get("aria-hidden", "")).lower() == "true":
        return True
    return False


def _extract_html_texts(raw: bytes, *, max_nodes: int = 200_000) -> Tuple[str, str]:
    try:
        from bs4 import BeautifulSoup
    except Exception as exc:  # pragma: no cover - dependency gate
        raise _missing_dependency_error("HTML", "beautifulsoup4") from exc

    parser = "lxml"
    try:
        __import__("lxml")
    except Exception:
        parser = "html.parser"

    html = raw.decode("utf-8", errors="ignore")
    soup = BeautifulSoup(html, parser)
    node_count = sum(1 for _ in soup.descendants)
    if node_count > int(max_nodes):
        raise ValueError(f"html node limit exceeded: {node_count}>{max_nodes}")

    hidden_texts: List[str] = []
    hidden_tags = [t for t in soup.find_all(True) if _element_is_hidden(t)]
    for tag in hidden_tags:
        txt = _normalize_text(tag.get_text(" ", strip=True))
        if txt:
            hidden_texts.append(txt)
        tag.decompose()

    visible = _normalize_text(soup.get_text(" ", strip=True))
    hidden = _normalize_text(" ".join(hidden_texts))
    return visible, hidden


def _image_suffix(*, filename: str | None, mime: str | None) -> str:
    lower_name = str(filename or "").strip().lower()
    if lower_name.endswith(".png") or str(mime or "").strip().lower() == "image/png":
        return ".png"
    if lower_name.endswith(".webp") or str(mime or "").strip().lower() == "image/webp":
        return ".webp"
    if lower_name.endswith(".gif") or str(mime or "").strip().lower() == "image/gif":
        return ".gif"
    return ".jpg"


def _detect_image_size(raw: bytes) -> Tuple[int, int]:
    try:
        from PIL import Image  # type: ignore

        with Image.open(io.BytesIO(raw)) as img:
            width, height = img.size
            return max(0, int(width)), max(0, int(height))
    except Exception:
        return 0, 0


_IMAGE_MIME_BY_FORMAT = {
    "PNG": "image/png",
    "JPEG": "image/jpeg",
    "WEBP": "image/webp",
    "GIF": "image/gif",
}


def _make_visual_asset(
    raw: bytes,
    *,
    asset_id: str,
    role: str,
    source_kind: str,
    cfg: AttachmentVisualConfig,
    page_number: Optional[int] = None,
    embedded_index: Optional[int] = None,
) -> Optional[AttachmentVisualAsset]:
    if not raw or len(raw) > int(cfg.max_asset_bytes):
        return None
    try:
        from PIL import Image

        with Image.open(io.BytesIO(raw)) as image:
            image.load()
            width, height = image.size
            if width < int(cfg.min_width) or height < int(cfg.min_height):
                return None
            fmt = str(image.format or "").upper()
            mime = _IMAGE_MIME_BY_FORMAT.get(fmt)
            payload = bytes(raw)
            if mime is None:
                out = io.BytesIO()
                image.convert("RGB").save(out, format="PNG", optimize=True)
                payload = out.getvalue()
                mime = "image/png"
            if len(payload) > int(cfg.max_asset_bytes):
                return None
    except Exception:
        return None
    digest = hashlib.sha256(payload).hexdigest()
    return AttachmentVisualAsset(
        asset_id=str(asset_id),
        mime=str(mime),
        sha256=digest,
        payload_b64=base64.b64encode(payload).decode("ascii"),
        role=str(role),
        source_kind=str(source_kind),
        page_number=page_number,
        embedded_index=embedded_index,
        width=int(width),
        height=int(height),
        size_bytes=len(payload),
    )


def _bounded_visual_assets(
    rows: Sequence[AttachmentVisualAsset], cfg: AttachmentVisualConfig
) -> List[AttachmentVisualAsset]:
    out: List[AttachmentVisualAsset] = []
    total = 0
    seen: set[str] = set()
    for row in rows:
        if row.sha256 in seen:
            continue
        if len(out) >= int(cfg.max_assets):
            break
        if total + int(row.size_bytes) > int(cfg.max_total_bytes):
            break
        seen.add(row.sha256)
        total += int(row.size_bytes)
        out.append(row)
    return out


def _render_pdf_visual_pages(
    raw: bytes, *, cfg: AttachmentVisualConfig
) -> List[AttachmentVisualAsset]:
    if not cfg.enabled or not cfg.render_pdf_pages:
        return []
    try:
        import fitz  # PyMuPDF
    except Exception as exc:
        raise _missing_dependency_error("Visual PDF", "PyMuPDF") from exc
    doc = fitz.open(stream=raw, filetype="pdf")
    rows: List[AttachmentVisualAsset] = []
    scale = float(cfg.pdf_dpi) / 72.0
    try:
        for idx in range(min(len(doc), int(cfg.max_pdf_pages))):
            page = doc.load_page(idx)
            rect = page.rect
            requested_pixels = max(1.0, float(rect.width) * scale) * max(
                1.0, float(rect.height) * scale
            )
            page_scale = scale
            if requested_pixels > float(cfg.max_pdf_pixels_per_page):
                page_scale *= (
                    float(cfg.max_pdf_pixels_per_page) / requested_pixels
                ) ** 0.5
            pix = page.get_pixmap(
                matrix=fitz.Matrix(page_scale, page_scale), alpha=False
            )
            actual_pixels = int(pix.width) * int(pix.height)
            if actual_pixels > int(cfg.max_pdf_pixels_per_page):
                # Rasterizers round dimensions up; correct once from observed size.
                page_scale *= (
                    float(cfg.max_pdf_pixels_per_page) / float(actual_pixels)
                ) ** 0.5 * 0.995
                pix = page.get_pixmap(
                    matrix=fitz.Matrix(page_scale, page_scale), alpha=False
                )
                actual_pixels = int(pix.width) * int(pix.height)
            if actual_pixels > int(cfg.max_pdf_pixels_per_page):
                raise ValueError("visual pdf page pixel limit exceeded")
            payload = pix.tobytes("png")
            asset = _make_visual_asset(
                payload,
                asset_id=f"pdf-page-{idx + 1}",
                role="full_page_context",
                source_kind="pdf_page",
                page_number=idx + 1,
                cfg=cfg,
            )
            if asset is not None:
                rows.append(asset)
    finally:
        doc.close()
    return _bounded_visual_assets(rows, cfg)


def _extract_docx_visual_assets(
    raw: bytes, *, cfg: AttachmentVisualConfig
) -> List[AttachmentVisualAsset]:
    if not cfg.enabled or not cfg.extract_embedded_images:
        return []
    rows: List[AttachmentVisualAsset] = []
    with zipfile.ZipFile(io.BytesIO(raw)) as zf:
        names = sorted(
            name for name in zf.namelist() if name.lower().startswith("word/media/")
        )
        for idx, name in enumerate(names):
            if len(rows) >= int(cfg.max_assets):
                break
            info = zf.getinfo(name)
            if int(info.file_size) > int(cfg.max_asset_bytes):
                continue
            asset = _make_visual_asset(
                zf.read(name),
                asset_id=f"docx-embedded-{idx + 1}",
                role="untrusted_visual_content",
                source_kind="docx_embedded",
                embedded_index=idx + 1,
                cfg=cfg,
            )
            if asset is not None:
                rows.append(asset)
    return _bounded_visual_assets(rows, cfg)


def _extract_html_visual_assets(
    raw: bytes, *, cfg: AttachmentVisualConfig
) -> List[AttachmentVisualAsset]:
    if not cfg.enabled or not cfg.extract_embedded_images:
        return []
    try:
        from bs4 import BeautifulSoup
    except Exception as exc:
        raise _missing_dependency_error("HTML visual", "beautifulsoup4") from exc
    html = raw.decode("utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")
    rows: List[AttachmentVisualAsset] = []
    for idx, tag in enumerate(soup.find_all("img")):
        if len(rows) >= int(cfg.max_assets):
            break
        src = str(tag.attrs.get("src", "") or "").strip()
        if not src.lower().startswith("data:image/") or ";base64," not in src.lower():
            continue  # Remote fetching is intentionally forbidden.
        header, encoded = src.split(",", 1)
        # Reject oversized data URIs before allocating decoded bytes.
        if len(encoded) > ((int(cfg.max_asset_bytes) + 2) // 3) * 4 + 8:
            continue
        try:
            payload = base64.b64decode(encoded, validate=True)
        except Exception:
            continue
        asset = _make_visual_asset(
            payload,
            asset_id=f"html-embedded-{idx + 1}",
            role="untrusted_visual_content",
            source_kind="html_data_uri",
            embedded_index=idx + 1,
            cfg=cfg,
        )
        if asset is not None:
            rows.append(asset)
    return _bounded_visual_assets(rows, cfg)


def _coerce_polygon_input(raw_value: Any) -> Any:
    if isinstance(raw_value, (list, tuple)):
        return raw_value
    tolist = getattr(raw_value, "tolist", None)
    if callable(tolist):
        try:
            return tolist()
        except Exception:
            return None
    return raw_value


def _coerce_polygon_px(
    raw_bbox: Any,
    *,
    image_width: int,
    image_height: int,
) -> Optional[Tuple[Tuple[float, float], ...]]:
    raw_bbox = _coerce_polygon_input(raw_bbox)
    if not isinstance(raw_bbox, (list, tuple)):
        return None
    out: List[Tuple[float, float]] = []
    for point in list(raw_bbox):
        if not isinstance(point, (list, tuple)) or len(point) < 2:
            return None
        try:
            x = float(point[0])
            y = float(point[1])
        except (TypeError, ValueError):
            return None
        if not (x == x and y == y):  # NaN
            return None
        if x in {float("inf"), float("-inf")} or y in {float("inf"), float("-inf")}:
            return None
        if x < 0.0 or y < 0.0:
            return None
        if (
            int(image_width) > 0
            and int(image_height) > 0
            and (x > float(image_width) or y > float(image_height))
        ):
            return None
        out.append((x, y))
    if len(out) < 3:
        return None
    xs = [p[0] for p in out]
    ys = [p[1] for p in out]
    if (max(xs) - min(xs)) <= 0.0 or (max(ys) - min(ys)) <= 0.0:
        return None
    return tuple(out)


def _clip_confidence(value: Any) -> Optional[float]:
    try:
        out = float(value)
    except (TypeError, ValueError):
        return None
    return max(0.0, min(1.0, out))


def _normalize_ocr_spans(spans: Sequence[OCRSpan], *, max_chars: int) -> List[OCRSpan]:
    out: List[OCRSpan] = []
    total = 0
    for idx, span in enumerate(list(spans or [])):
        text = _normalize_text(str(span.text or ""))
        if not text:
            continue
        if total >= max_chars:
            break
        remaining = max(0, max_chars - total)
        clipped_text = text[:remaining].strip() if len(text) > remaining else text
        if not clipped_text:
            break
        char_start = total if not out else total + 1
        char_end = char_start + len(clipped_text)
        out.append(
            OCRSpan(
                span_id=str(span.span_id or f"ocr-span-{idx:04d}"),
                text=clipped_text,
                confidence=_clip_confidence(span.confidence),
                polygon_px=span.polygon_px,
                image_width=max(0, int(span.image_width or 0)),
                image_height=max(0, int(span.image_height or 0)),
                provider_order=int(span.provider_order)
                if span.provider_order is not None
                else idx,
                char_start=char_start,
                char_end=char_end,
            )
        )
        total = char_end
        if len(text) > remaining:
            break
    return out


def _spans_to_text(spans: Sequence[OCRSpan]) -> str:
    return _normalize_text(
        " ".join(_normalize_text(str(span.text or "")) for span in list(spans or []))
    )


def _parse_paddleocr_lines(raw_result: Any) -> List[OCRSpan]:
    spans: List[OCRSpan] = []
    stack: List[Any] = [raw_result]
    order = 0
    while stack:
        current = stack.pop()
        if isinstance(current, tuple):
            stack.extend(list(current))
            continue
        if isinstance(current, list):
            if len(current) == 2 and isinstance(current[1], (list, tuple)):
                maybe_text = current[1][0] if current[1] else None
                if isinstance(maybe_text, str):
                    norm = _normalize_text(maybe_text)
                    if norm:
                        polygon_px = _coerce_polygon_input(current[0])
                        spans.append(
                            OCRSpan(
                                span_id=f"ocr-span-{order:04d}",
                                text=norm,
                                confidence=_clip_confidence(
                                    current[1][1] if len(current[1]) > 1 else None
                                ),
                                polygon_px=polygon_px,
                                provider_order=order,
                            )
                        )
                        order += 1
                    continue
            stack.extend(reversed(current))
    return spans


def _parse_rapidocr_lines(raw_result: Any) -> List[OCRSpan]:
    spans: List[OCRSpan] = []
    txts = _coerce_polygon_input(getattr(raw_result, "txts", None))
    boxes = _coerce_polygon_input(getattr(raw_result, "boxes", None))
    scores = _coerce_polygon_input(getattr(raw_result, "scores", None))
    if isinstance(txts, (list, tuple)):
        for idx, item in enumerate(txts):
            norm = _normalize_text(str(item))
            if norm:
                polygon_px = None
                confidence = None
                if isinstance(boxes, (list, tuple)) and idx < len(boxes):
                    polygon_px = _coerce_polygon_input(boxes[idx])
                if isinstance(scores, (list, tuple)) and idx < len(scores):
                    confidence = _clip_confidence(scores[idx])
                spans.append(
                    OCRSpan(
                        span_id=f"ocr-span-{idx:04d}",
                        text=norm,
                        confidence=confidence,
                        polygon_px=polygon_px,  # normalized later once image dims are known
                        provider_order=idx,
                    )
                )
    if spans:
        return spans
    word_results = getattr(raw_result, "word_results", None)
    if isinstance(word_results, (list, tuple)):
        for idx, item in enumerate(word_results):
            if not isinstance(item, (list, tuple)) or not item:
                continue
            maybe_text = item[0]
            if isinstance(maybe_text, str):
                norm = _normalize_text(maybe_text)
                if norm:
                    confidence = _clip_confidence(item[1] if len(item) > 1 else None)
                    spans.append(
                        OCRSpan(
                            span_id=f"ocr-span-{idx:04d}",
                            text=norm,
                            confidence=confidence,
                            polygon_px=(item[2] if len(item) > 2 else None),
                            provider_order=idx,
                        )
                    )
    return spans


def _bind_image_geometry(
    spans: Sequence[OCRSpan], *, image_width: int, image_height: int
) -> List[OCRSpan]:
    out: List[OCRSpan] = []
    for idx, span in enumerate(list(spans or [])):
        width = max(int(image_width), int(getattr(span, "image_width", 0) or 0))
        height = max(int(image_height), int(getattr(span, "image_height", 0) or 0))
        polygon_px = _coerce_polygon_px(
            span.polygon_px,
            image_width=width,
            image_height=height,
        )
        out.append(
            OCRSpan(
                span_id=str(span.span_id or f"ocr-span-{idx:04d}"),
                text=str(span.text or ""),
                confidence=_clip_confidence(span.confidence),
                polygon_px=polygon_px,
                image_width=max(0, width),
                image_height=max(0, height),
                provider_order=(
                    int(span.provider_order) if span.provider_order is not None else idx
                ),
                char_start=span.char_start,
                char_end=span.char_end,
            )
        )
    return out


def _build_ocr_chunks_from_spans(
    spans: Sequence[OCRSpan],
    *,
    max_chunk_chars: int,
    chunk_overlap: int,
    max_spans_per_chunk: int,
) -> Tuple[str, List[AttachmentChunk]]:
    normalized_spans = list(spans or [])
    full_text = _spans_to_text(normalized_spans)
    if not full_text:
        return "", []
    chunks: List[AttachmentChunk] = []
    idx = 0
    overlap_spans = 1 if int(chunk_overlap) > 0 else 0
    span_cap = max(1, int(max_spans_per_chunk))
    while idx < len(normalized_spans):
        chunk_spans: List[OCRSpan] = []
        chunk_chars = 0
        while (
            idx + len(chunk_spans) < len(normalized_spans)
            and len(chunk_spans) < span_cap
        ):
            span = normalized_spans[idx + len(chunk_spans)]
            span_text = _normalize_text(str(span.text or ""))
            extra_chars = len(span_text) + (1 if chunk_spans else 0)
            if chunk_spans and (chunk_chars + extra_chars) > int(max_chunk_chars):
                break
            chunk_spans.append(span)
            chunk_chars += extra_chars
        if not chunk_spans:
            chunk_spans = [normalized_spans[idx]]
        chunk_text = _normalize_text(
            " ".join(str(span.text or "") for span in chunk_spans)
        )
        chunk_start = min(int(span.char_start or 0) for span in chunk_spans)
        chunk_end = max(
            int(span.char_end or int(span.char_start or 0)) for span in chunk_spans
        )
        span_ids = [str(span.span_id) for span in chunk_spans]
        chunks.append(
            AttachmentChunk(
                text=chunk_text,
                kind="ocr",
                is_hidden=False,
                ocr_span_ids=span_ids,
                char_start=int(chunk_start),
                char_end=int(chunk_end),
            )
        )
        idx += max(1, len(chunk_spans) - overlap_spans)
    return full_text, chunks


def _get_paddleocr_engine(cfg: AttachmentOCRConfig) -> Any:
    from paddleocr import PaddleOCR  # type: ignore

    cache_key = (str(cfg.lang or "en"), bool(cfg.use_angle_cls))
    with _PADDLE_OCR_CACHE_LOCK:
        cached = _PADDLE_OCR_CACHE.get(cache_key)
        if cached is not None:
            return cached
        engine = PaddleOCR(
            use_angle_cls=bool(cfg.use_angle_cls), lang=str(cfg.lang or "en")
        )
        _PADDLE_OCR_CACHE[cache_key] = engine
        return engine


def _get_rapidocr_engine(cfg: AttachmentOCRConfig) -> Any:
    from rapidocr import RapidOCR  # type: ignore

    cache_key = (str(cfg.lang or "en"),)
    with _RAPIDOCR_CACHE_LOCK:
        cached = _RAPIDOCR_CACHE.get(cache_key)
        if cached is not None:
            return cached
        # RapidOCR is primarily multilingual without per-call language switching in the same way
        # as PaddleOCR, so we keep the runtime contract stable and cache one engine per lang hint.
        engine = RapidOCR()
        _RAPIDOCR_CACHE[cache_key] = engine
        return engine


def _extract_image_ocr_text_from_bytes(
    raw: bytes,
    *,
    filename: str | None,
    mime: str | None,
    cfg: AttachmentOCRConfig,
) -> ImageOCRResult:
    enabled = str(cfg.enabled).strip().lower()
    provider = str(cfg.provider).strip().lower() or "rapidocr"
    if enabled == "false":
        return ImageOCRResult(
            text="", status="disabled", provider=provider, warnings=["ocr_disabled"]
        )
    if provider not in {"paddleocr", "rapidocr"}:
        return ImageOCRResult(
            text="",
            status="unavailable",
            provider=provider,
            warnings=["ocr_unavailable"],
        )

    try:
        if provider == "rapidocr":
            suffix = _image_suffix(filename=filename, mime=mime)
            if str(cfg.execution_mode).strip().lower() == "persistent_worker":
                from omega.vision.ocr_runtime import (
                    OCRWorkerSettings,
                    recognize_with_worker,
                )

                spans = recognize_with_worker(
                    raw,
                    suffix=suffix,
                    use_angle_cls=bool(cfg.use_angle_cls),
                    settings=OCRWorkerSettings(
                        provider="rapidocr",
                        startup_timeout_sec=float(cfg.worker_startup_timeout_sec),
                        request_timeout_sec=float(cfg.worker_request_timeout_sec),
                        max_memory_mb=int(cfg.worker_max_memory_mb),
                        max_requests_per_worker=int(cfg.worker_max_requests),
                        pool_size=int(cfg.worker_pool_size),
                        max_pending_requests=int(cfg.worker_max_pending_requests),
                        queue_timeout_sec=float(cfg.worker_queue_timeout_sec),
                        intra_op_num_threads=int(cfg.worker_intra_op_threads),
                        inter_op_num_threads=int(cfg.worker_inter_op_threads),
                    ),
                )
            else:
                ocr = _get_rapidocr_engine(cfg)
                tmp_path = ""
                try:
                    with tempfile.NamedTemporaryFile(
                        delete=False, suffix=suffix
                    ) as tmp:
                        tmp.write(raw)
                        tmp.flush()
                        tmp_path = str(tmp.name)
                    result = ocr(tmp_path, use_cls=bool(cfg.use_angle_cls))
                finally:
                    if tmp_path:
                        try:
                            Path(tmp_path).unlink(missing_ok=True)
                        except Exception:
                            pass
                spans = _parse_rapidocr_lines(result)
        else:
            ocr = _get_paddleocr_engine(cfg)
            suffix = _image_suffix(filename=filename, mime=mime)
            with tempfile.NamedTemporaryFile(delete=True, suffix=suffix) as tmp:
                tmp.write(raw)
                tmp.flush()
                if hasattr(ocr, "predict"):
                    result = ocr.predict(tmp.name)
                else:
                    result = ocr.ocr(tmp.name, cls=bool(cfg.use_angle_cls))
            spans = _parse_paddleocr_lines(result)
        clipped_spans = _normalize_ocr_spans(
            spans, max_chars=max(1, int(cfg.max_text_chars))
        )
        full_text = _spans_to_text(clipped_spans)
        if not full_text:
            return ImageOCRResult(
                text="",
                status="empty",
                provider=provider,
                warnings=["ocr_empty"],
                spans=[],
            )
        return ImageOCRResult(
            text=full_text,
            status="success",
            provider=provider,
            warnings=["ocr_text_present"],
            spans=clipped_spans,
        )
    except ModuleNotFoundError:
        return ImageOCRResult(
            text="",
            status="unavailable",
            provider=provider,
            warnings=["ocr_unavailable"],
            spans=[],
        )
    except OCRWorkerOverloadedError:
        return ImageOCRResult(
            text="",
            status="overloaded",
            provider=provider,
            warnings=["ocr_overloaded"],
            spans=[],
        )
    except TimeoutError:
        return ImageOCRResult(
            text="",
            status="timeout",
            provider=provider,
            warnings=["ocr_timeout"],
            spans=[],
        )
    except MemoryError:
        return ImageOCRResult(
            text="",
            status="resource_exceeded",
            provider=provider,
            warnings=["ocr_resource_exceeded"],
            spans=[],
        )
    except Exception:
        return ImageOCRResult(
            text="", status="error", provider=provider, warnings=["ocr_error"], spans=[]
        )


def _is_scan_like_pdf(
    text: str, pages_count: int, cfg: AttachmentIngestionConfig
) -> bool:
    if pages_count <= 0:
        return False
    text_norm = _normalize_text(text)
    if not text_norm:
        return True
    chars_per_page = len(text_norm) / float(pages_count)
    alpha_chars = sum(1 for ch in text_norm if ch.isalpha())
    alpha_ratio = float(alpha_chars) / float(max(1, len(text_norm)))
    return chars_per_page < float(
        cfg.scan_like_min_chars_per_page
    ) or alpha_ratio < float(cfg.scan_like_min_alpha_ratio)


def _sniff_format(raw: bytes) -> str:
    head = bytes(raw[:64])
    low = head.lstrip().lower()
    if head.startswith(b"%PDF-"):
        return "pdf"
    if head.startswith(b"PK\x03\x04") or head.startswith(b"PK\x05\x06"):
        try:
            with zipfile.ZipFile(io.BytesIO(raw)) as zf:
                names = set(zf.namelist())
            if "[Content_Types].xml" in names and "word/document.xml" in names:
                return "docx"
        except Exception:
            pass
        return "zip"
    if (
        head.startswith(b"\x89PNG\r\n\x1a\n")
        or head[:3] == b"\xff\xd8\xff"
        or head.startswith((b"GIF87a", b"GIF89a"))
        or (head.startswith(b"RIFF") and b"WEBP" in head[:16])
    ):
        return "image"
    if low.startswith((b"<!doctype html", b"<html", b"<head", b"<body")):
        return "html"
    return "text"


def _preflight_attachment(
    raw: bytes, *, declared_format: str, cfg: AttachmentIngestionConfig
) -> str:
    if not cfg.strict_magic and not cfg.sandbox.enabled:
        return declared_format
    sniffed = _sniff_format(raw)
    structured = {"pdf", "docx", "html", "image", "zip"}
    if (
        cfg.strict_magic
        and declared_format in structured
        and sniffed != declared_format
    ):
        raise ValueError(
            f"attachment format mismatch: declared={declared_format} detected={sniffed}"
        )
    fmt = sniffed if sniffed in structured else declared_format
    if fmt == "docx":
        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            infos = zf.infolist()
            if len(infos) > cfg.max_docx_entries:
                raise ValueError("docx entry limit exceeded")
            total = 0
            for info in infos:
                name = str(info.filename)
                parts = Path(name).parts
                if name.startswith(("/", "\\")) or any(
                    part in {"..", ""} for part in parts
                ):
                    raise ValueError("docx contains unsafe archive path")
                if info.flag_bits & 0x1:
                    raise ValueError("encrypted docx entries are forbidden")
                total += int(info.file_size)
                if total > cfg.max_docx_uncompressed_bytes:
                    raise ValueError("docx uncompressed size limit exceeded")
    elif fmt == "pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(raw), strict=True)
            if len(reader.pages) > cfg.max_pdf_pages:
                raise ValueError("pdf page limit exceeded")
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError("invalid pdf") from exc
    elif fmt == "image":
        try:
            from PIL import Image

            with Image.open(io.BytesIO(raw)) as image:
                width, height = image.size
                if int(width) * int(height) > cfg.max_image_pixels:
                    raise ValueError("image pixel limit exceeded")
                image.verify()
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError("invalid image") from exc
    return fmt


def _sandbox_result_from_dict(payload: Mapping[str, Any]) -> AttachmentExtractResult:
    chunks = [
        AttachmentChunk(
            text=str(row.get("text", "")),
            kind=str(row.get("kind", "")),
            is_hidden=bool(row.get("is_hidden", False)),
            ocr_span_ids=[
                str(value) for value in list(row.get("ocr_span_ids", []) or [])
            ],
            char_start=(
                int(row["char_start"]) if row.get("char_start") is not None else None
            ),
            char_end=(
                int(row["char_end"]) if row.get("char_end") is not None else None
            ),
        )
        for row in list(payload.get("chunks", []) or [])
        if isinstance(row, Mapping)
    ]
    spans: List[OCRSpan] = []
    for row in list(payload.get("ocr_spans", []) or []):
        if not isinstance(row, Mapping):
            continue
        polygon_raw = row.get("polygon_px")
        polygon = None
        if polygon_raw is not None:
            polygon = tuple(
                (float(point[0]), float(point[1])) for point in list(polygon_raw)
            )
        spans.append(
            OCRSpan(
                span_id=str(row.get("span_id", "")),
                text=str(row.get("text", "")),
                confidence=(
                    float(row["confidence"])
                    if row.get("confidence") is not None
                    else None
                ),
                polygon_px=polygon,
                image_width=int(row.get("image_width", 0)),
                image_height=int(row.get("image_height", 0)),
                provider_order=(
                    int(row["provider_order"])
                    if row.get("provider_order") is not None
                    else None
                ),
                char_start=(
                    int(row["char_start"])
                    if row.get("char_start") is not None
                    else None
                ),
                char_end=(
                    int(row["char_end"]) if row.get("char_end") is not None else None
                ),
            )
        )
    visual_assets: List[AttachmentVisualAsset] = []
    for row in list(payload.get("visual_assets", []) or []):
        if not isinstance(row, Mapping):
            continue
        visual_assets.append(
            AttachmentVisualAsset(
                asset_id=str(row.get("asset_id", "")),
                mime=str(row.get("mime", "")),
                sha256=str(row.get("sha256", "")),
                payload_b64=str(row.get("payload_b64", "")),
                role=str(row.get("role", "untrusted_visual_content")),
                source_kind=str(row.get("source_kind", "embedded")),
                page_number=(
                    int(row["page_number"])
                    if row.get("page_number") is not None
                    else None
                ),
                embedded_index=(
                    int(row["embedded_index"])
                    if row.get("embedded_index") is not None
                    else None
                ),
                width=int(row.get("width", 0)),
                height=int(row.get("height", 0)),
                size_bytes=int(row.get("size_bytes", 0)),
            )
        )
    return AttachmentExtractResult(
        text=str(payload.get("text", "")),
        chunks=chunks,
        format=str(payload.get("format", "")),
        text_empty=bool(payload.get("text_empty", False)),
        scan_like=bool(payload.get("scan_like", False)),
        hidden_text_chars=int(payload.get("hidden_text_chars", 0)),
        warnings=[str(value) for value in list(payload.get("warnings", []) or [])],
        recommended_verdict=str(payload.get("recommended_verdict", "quarantine")),
        is_image=bool(payload.get("is_image", False)),
        image_mime=(
            str(payload["image_mime"])
            if payload.get("image_mime") is not None
            else None
        ),
        image_sha256=(
            str(payload["image_sha256"])
            if payload.get("image_sha256") is not None
            else None
        ),
        image_bytes_size=int(payload.get("image_bytes_size", 0)),
        ocr_status=str(payload.get("ocr_status", "none")),
        ocr_provider=(
            str(payload["ocr_provider"])
            if payload.get("ocr_provider") is not None
            else None
        ),
        ocr_text_chars=int(payload.get("ocr_text_chars", 0)),
        ocr_spans=spans,
        ocr_quality=OCRQualitySummary(**dict(payload.get("ocr_quality", {}) or {})),
        visual_assets=visual_assets,
        visual_status=str(payload.get("visual_status", "none")),
    )


def _sandbox_environment() -> Dict[str, str]:
    # Preserve only interpreter/runtime essentials. Provider credentials and other
    # application secrets are intentionally not inherited by untrusted parsers.
    allowed = {
        "PATH",
        "LANG",
        "LC_ALL",
        "TZ",
        "SYSTEMROOT",
        "WINDIR",
        "TEMP",
        "TMP",
        "TMPDIR",
        "HOME",
    }
    env = {key: value for key, value in os.environ.items() if key in allowed}
    package_root = str(Path(__file__).resolve().parents[2])
    env["PYTHONPATH"] = package_root
    env["PYTHONNOUSERSITE"] = "1"
    env["PYTHONDONTWRITEBYTECODE"] = "1"
    return env


def _terminate_parser_process(proc: subprocess.Popen[bytes]) -> None:
    if proc.poll() is not None:
        return
    try:
        if os.name == "posix":
            os.killpg(proc.pid, signal.SIGKILL)
        else:
            proc.kill()
    except ProcessLookupError:
        pass
    finally:
        try:
            proc.wait(timeout=1.0)
        except subprocess.TimeoutExpired:
            pass


def _extract_attachment_in_sandbox(
    *,
    raw: bytes,
    filename: str | None,
    mime: str | None,
    cfg: Mapping[str, Any] | None,
    parsed_cfg: AttachmentIngestionConfig,
) -> AttachmentExtractResult:
    with tempfile.TemporaryDirectory(prefix="omega-attachment-") as tmp_raw:
        tmp = Path(tmp_raw).resolve()
        input_path = tmp / "input.bin"
        request_path = tmp / "request.json"
        response_path = tmp / "response.json"
        input_path.write_bytes(raw)
        input_path.chmod(0o600)
        request_payload = {
            "input_path": str(input_path),
            "filename": filename,
            "mime": mime,
            "cfg": dict(cfg or {}),
            "max_memory_mb": int(parsed_cfg.sandbox.max_memory_mb),
            "max_cpu_sec": int(parsed_cfg.sandbox.max_cpu_sec),
        }
        request_path.write_text(
            json.dumps(
                request_payload,
                ensure_ascii=False,
                allow_nan=False,
                separators=(",", ":"),
            ),
            encoding="utf-8",
        )
        request_path.chmod(0o600)
        worker_argv = [
            sys.executable,
            str(Path(__file__).with_name("attachment_sandbox_worker.py")),
            str(request_path),
            str(response_path),
        ]
        if os.name == "posix":
            from omega.rag.attachment_parser_runtime import run_attachment_parser

            return_code = run_attachment_parser(
                request_path=request_path,
                response_path=response_path,
                timeout_sec=float(parsed_cfg.sandbox.timeout_sec),
            )
        else:
            proc = subprocess.Popen(
                worker_argv,
                stdin=subprocess.DEVNULL,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                cwd=str(tmp),
                env=_sandbox_environment(),
                close_fds=True,
                start_new_session=True,
            )
            try:
                return_code = proc.wait(timeout=float(parsed_cfg.sandbox.timeout_sec))
            except subprocess.TimeoutExpired as exc:
                _terminate_parser_process(proc)
                raise TimeoutError("attachment parser sandbox timeout") from exc
            finally:
                if proc.poll() is None:
                    _terminate_parser_process(proc)

        if not response_path.exists():
            raise ValueError(
                f"attachment parser sandbox exited without response: code={return_code}"
            )
        visual_transport_budget = (
            int(parsed_cfg.visual.max_total_bytes * 1.40)
            if parsed_cfg.visual.enabled
            else 0
        )
        max_response_bytes = max(
            1_048_576,
            int(parsed_cfg.max_extracted_chars) * 12
            + 1_048_576
            + visual_transport_budget,
        )
        if response_path.stat().st_size > max_response_bytes:
            raise ValueError("attachment parser sandbox response exceeds limit")
        response = json.loads(response_path.read_text(encoding="utf-8"))
        if not isinstance(response, Mapping) or not bool(response.get("ok", False)):
            error_type = (
                str(response.get("error_type", "ParserError"))
                if isinstance(response, Mapping)
                else "ParserError"
            )
            error = (
                str(response.get("error", "unknown parser error"))
                if isinstance(response, Mapping)
                else "invalid response"
            )
            raise ValueError(f"attachment parser sandbox failed: {error_type}: {error}")
        result = response.get("result")
        if not isinstance(result, Mapping):
            raise ValueError("attachment parser sandbox returned invalid result")
        return _sandbox_result_from_dict(result)


def _extract_attachment_in_process(
    *,
    path: str | Path | None = None,
    content_bytes: bytes | None = None,
    filename: str | None = None,
    mime: str | None = None,
    cfg: Mapping[str, Any] | None = None,
) -> AttachmentExtractResult:
    parsed_cfg = AttachmentIngestionConfig.from_cfg(cfg)
    if not parsed_cfg.enabled:
        return AttachmentExtractResult(
            text="",
            chunks=[],
            format="disabled",
            text_empty=True,
            scan_like=False,
            hidden_text_chars=0,
            warnings=["attachments_disabled"],
            recommended_verdict="allow",
            is_image=False,
            image_mime=None,
            image_sha256=None,
            image_bytes_size=0,
        )

    path_obj = Path(path) if path is not None else None
    fmt = _detect_format(path=path_obj, filename=filename, mime=mime)
    raw = _load_raw_bytes(path=path_obj, content_bytes=content_bytes)
    warnings: List[str] = []

    if len(raw) > parsed_cfg.max_file_bytes:
        raise ValueError(
            f"attachment exceeds max_file_bytes={parsed_cfg.max_file_bytes}: got={len(raw)} format={fmt}"
        )

    if fmt == "zip":
        warnings.append("zip_deferred_runtime")
        return AttachmentExtractResult(
            text="",
            chunks=[],
            format="zip",
            text_empty=True,
            scan_like=False,
            hidden_text_chars=0,
            warnings=warnings,
            recommended_verdict="quarantine",
            is_image=False,
            image_mime=None,
            image_sha256=None,
            image_bytes_size=0,
        )

    text = ""
    scan_like = False
    hidden_text = ""
    is_image = False
    image_mime: Optional[str] = None
    image_sha256: Optional[str] = None
    ocr_status = "none"
    ocr_provider: Optional[str] = None
    ocr_text_chars = 0
    ocr_spans: List[OCRSpan] = []
    ocr_quality = OCRQualitySummary()
    visual_assets: List[AttachmentVisualAsset] = []
    visual_status = "disabled" if not parsed_cfg.visual.enabled else "empty"

    if fmt == "pdf":
        text, pages = _extract_pdf_text_from_bytes(raw)
        scan_like = _is_scan_like_pdf(text=text, pages_count=pages, cfg=parsed_cfg)
        if parsed_cfg.visual.enabled:
            try:
                visual_assets = _render_pdf_visual_pages(raw, cfg=parsed_cfg.visual)
                visual_status = "success" if visual_assets else "empty"
            except Exception:
                visual_status = "error"
                warnings.append("visual_pdf_render_error")
    elif fmt == "docx":
        text = _extract_docx_text_from_bytes(raw)
        if parsed_cfg.visual.enabled:
            try:
                visual_assets = _extract_docx_visual_assets(raw, cfg=parsed_cfg.visual)
                visual_status = "success" if visual_assets else "empty"
            except Exception:
                visual_status = "error"
                warnings.append("visual_embedded_extract_error")
    elif fmt == "html":
        visible, hidden = _extract_html_texts(raw, max_nodes=parsed_cfg.max_html_nodes)
        text = visible
        hidden_text = hidden if parsed_cfg.html_include_hidden else ""
        if parsed_cfg.visual.enabled:
            try:
                visual_assets = _extract_html_visual_assets(raw, cfg=parsed_cfg.visual)
                visual_status = "success" if visual_assets else "empty"
            except Exception:
                visual_status = "error"
                warnings.append("visual_embedded_extract_error")
    elif fmt == "image":
        is_image = True
        image_mime = str(mime or "").strip().lower() or "application/octet-stream"
        image_sha256 = _sha256_hex(raw)
        image_width, image_height = _detect_image_size(raw)
        if parsed_cfg.visual.enabled:
            asset = _make_visual_asset(
                raw,
                asset_id="image-original-1",
                role="untrusted_visual_content",
                source_kind="image_attachment",
                embedded_index=1,
                cfg=parsed_cfg.visual,
            )
            visual_assets = [asset] if asset is not None else []
            visual_status = "success" if visual_assets else "empty"
        try:
            ocr_result = _extract_image_ocr_text_from_bytes(
                raw,
                filename=filename,
                mime=image_mime,
                cfg=parsed_cfg.ocr,
            )
        except Exception:
            ocr_result = _ocr_error_result(parsed_cfg)
        text = str(ocr_result.text or "")
        ocr_status = str(ocr_result.status or "none")
        ocr_provider = str(ocr_result.provider) if ocr_result.provider else None
        ocr_text_chars = len(_normalize_text(text))
        ocr_spans, ocr_quality = normalize_ocr_spans(
            _bind_image_geometry(
                list(ocr_result.spans or []),
                image_width=int(image_width),
                image_height=int(image_height),
            ),
            max_chars=max(1, int(parsed_cfg.ocr.max_text_chars)),
            image_width=int(image_width),
            image_height=int(image_height),
            policy=OCRQualityPolicy(
                min_confidence=float(parsed_cfg.ocr.min_confidence),
                max_spans=int(parsed_cfg.ocr.max_spans),
                max_span_chars=int(parsed_cfg.ocr.max_span_chars),
                require_geometry=bool(parsed_cfg.ocr.require_geometry),
                min_polygon_area_px=float(parsed_cfg.ocr.min_polygon_area_px),
            ),
        )
        text = _spans_to_text(ocr_spans)
        ocr_text_chars = len(_normalize_text(text))
        warnings.extend(list(ocr_result.warnings or []))
        if str(ocr_result.status) == "success" and not ocr_spans:
            ocr_status = "filtered_empty"
            warnings.append("ocr_filtered_empty")
        if ocr_quality.dropped_low_confidence:
            warnings.append("ocr_low_confidence_spans_dropped")
        if ocr_quality.dropped_invalid_geometry:
            warnings.append("ocr_invalid_geometry_spans_dropped")
        if not text:
            warnings.append("image_semantic_only")
        if str(parsed_cfg.ocr.failure_policy) == "fail_closed" and str(ocr_status) in {
            "error",
            "unavailable",
            "timeout",
            "resource_exceeded",
            "overloaded",
        }:
            raise RuntimeError(f"ocr_required_{ocr_status}")
        if str(parsed_cfg.ocr.failure_policy) == "quarantine" and str(ocr_status) in {
            "error",
            "unavailable",
            "timeout",
            "resource_exceeded",
            "overloaded",
            "filtered_empty",
        }:
            warnings.append("ocr_failure_quarantine")
    else:
        text = raw.decode("utf-8", errors="ignore")

    if visual_assets:
        warnings.append("visual_assets_present")
    if parsed_cfg.visual.enabled and visual_status == "error":
        if parsed_cfg.visual.failure_policy == "fail_closed":
            raise RuntimeError("visual_extraction_required")
        if parsed_cfg.visual.failure_policy == "quarantine":
            warnings.append("visual_extraction_quarantine")

    text = _clip_text(text, max_chars=parsed_cfg.max_extracted_chars, warnings=warnings)
    hidden_text = _clip_text(
        hidden_text, max_chars=parsed_cfg.max_extracted_chars, warnings=warnings
    )

    primary_chunk_kind = (
        "ocr" if (bool(is_image) and str(ocr_status) == "success") else "visible"
    )
    chunks: List[AttachmentChunk] = []
    if primary_chunk_kind == "ocr":
        text, chunks = _build_ocr_chunks_from_spans(
            ocr_spans,
            max_chunk_chars=parsed_cfg.max_chunk_chars,
            chunk_overlap=parsed_cfg.chunk_overlap,
            max_spans_per_chunk=parsed_cfg.ocr.max_spans_per_chunk,
        )
    else:
        for c in _chunk_text(
            text,
            max_chunk_chars=parsed_cfg.max_chunk_chars,
            chunk_overlap=parsed_cfg.chunk_overlap,
        ):
            chunks.append(
                AttachmentChunk(text=c, kind=primary_chunk_kind, is_hidden=False)
            )

    hidden_chars = len(_normalize_text(hidden_text))
    if hidden_text:
        warnings.append("hidden_text_present")
        pref = parsed_cfg.hidden_chunk_prefix
        for c in _chunk_text(
            hidden_text,
            max_chunk_chars=parsed_cfg.max_chunk_chars,
            chunk_overlap=parsed_cfg.chunk_overlap,
        ):
            chunks.append(
                AttachmentChunk(
                    text=f"{pref}{c}".strip(), kind="hidden", is_hidden=True
                )
            )

    full_text = _normalize_text(" ".join(ch.text for ch in chunks))
    text_empty = len(full_text) == 0
    if text_empty:
        warnings.append("text_empty")
    if scan_like:
        warnings.append("scan_like")

    verdict = (
        "quarantine"
        if (
            text_empty
            or scan_like
            or "ocr_failure_quarantine" in warnings
            or "visual_extraction_quarantine" in warnings
        )
        else "allow"
    )
    return AttachmentExtractResult(
        text=full_text,
        chunks=chunks,
        format=fmt,
        text_empty=text_empty,
        scan_like=bool(scan_like),
        hidden_text_chars=hidden_chars,
        warnings=sorted(set(warnings)),
        recommended_verdict=verdict,
        is_image=bool(is_image),
        image_mime=image_mime,
        image_sha256=image_sha256,
        image_bytes_size=int(len(raw)) if is_image else 0,
        ocr_status=str(ocr_status),
        ocr_provider=ocr_provider,
        ocr_text_chars=int(ocr_text_chars),
        ocr_spans=list(ocr_spans),
        ocr_quality=ocr_quality,
        visual_assets=list(visual_assets),
        visual_status=str(visual_status),
    )


def _merge_image_ocr_result(
    *,
    base: AttachmentExtractResult,
    raw: bytes,
    ocr_result: ImageOCRResult,
    parsed_cfg: AttachmentIngestionConfig,
) -> AttachmentExtractResult:
    image_width, image_height = _detect_image_size(raw)
    spans, quality = normalize_ocr_spans(
        _bind_image_geometry(
            list(ocr_result.spans or []),
            image_width=int(image_width),
            image_height=int(image_height),
        ),
        max_chars=max(1, int(parsed_cfg.ocr.max_text_chars)),
        image_width=int(image_width),
        image_height=int(image_height),
        policy=OCRQualityPolicy(
            min_confidence=float(parsed_cfg.ocr.min_confidence),
            max_spans=int(parsed_cfg.ocr.max_spans),
            max_span_chars=int(parsed_cfg.ocr.max_span_chars),
            require_geometry=bool(parsed_cfg.ocr.require_geometry),
            min_polygon_area_px=float(parsed_cfg.ocr.min_polygon_area_px),
        ),
    )
    status = str(ocr_result.status or "none")
    if status == "success" and not spans:
        status = "filtered_empty"
    text, chunks = _build_ocr_chunks_from_spans(
        spans,
        max_chunk_chars=parsed_cfg.max_chunk_chars,
        chunk_overlap=parsed_cfg.chunk_overlap,
        max_spans_per_chunk=parsed_cfg.ocr.max_spans_per_chunk,
    )
    warnings = {str(value) for value in base.warnings}
    warnings.difference_update({"ocr_disabled", "image_semantic_only", "text_empty"})
    warnings.update(str(value) for value in list(ocr_result.warnings or []))
    if status == "filtered_empty":
        warnings.add("ocr_filtered_empty")
    if quality.dropped_low_confidence:
        warnings.add("ocr_low_confidence_spans_dropped")
    if quality.dropped_invalid_geometry:
        warnings.add("ocr_invalid_geometry_spans_dropped")
    if not text:
        warnings.update({"image_semantic_only", "text_empty"})
    failure_policy = str(parsed_cfg.ocr.failure_policy)
    if failure_policy == "fail_closed" and status in {
        "error",
        "unavailable",
        "timeout",
        "resource_exceeded",
        "overloaded",
    }:
        raise RuntimeError(f"ocr_required_{status}")
    if failure_policy == "quarantine" and status in {
        "error",
        "unavailable",
        "timeout",
        "resource_exceeded",
        "overloaded",
        "filtered_empty",
    }:
        warnings.add("ocr_failure_quarantine")
    verdict = "allow" if text else "quarantine"
    if "ocr_failure_quarantine" in warnings:
        verdict = "quarantine"
    return AttachmentExtractResult(
        text=text,
        chunks=chunks,
        format=base.format,
        text_empty=not bool(text),
        scan_like=base.scan_like,
        hidden_text_chars=base.hidden_text_chars,
        warnings=sorted(warnings),
        recommended_verdict=verdict,
        is_image=True,
        image_mime=base.image_mime,
        image_sha256=base.image_sha256,
        image_bytes_size=base.image_bytes_size,
        ocr_status=status,
        ocr_provider=ocr_result.provider,
        ocr_text_chars=len(_normalize_text(text)),
        ocr_spans=spans,
        ocr_quality=quality,
        visual_assets=list(base.visual_assets),
        visual_status=str(base.visual_status),
    )


def _ocr_error_result(parsed_cfg: AttachmentIngestionConfig) -> ImageOCRResult:
    return ImageOCRResult(
        text="",
        status="error",
        provider=str(parsed_cfg.ocr.provider),
        warnings=["ocr_error"],
    )


def _config_with_ocr_disabled(cfg: Mapping[str, Any] | None) -> Dict[str, Any]:
    import copy

    cloned: Dict[str, Any] = copy.deepcopy(dict(cfg or {}))
    ocr_cfg = cloned.setdefault("ocr", {})
    if not isinstance(ocr_cfg, dict):
        cloned["ocr"] = {}
        ocr_cfg = cloned["ocr"]
    ocr_cfg["enabled"] = "false"
    return cloned


def extract_attachment(
    *,
    path: str | Path | None = None,
    content_bytes: bytes | None = None,
    filename: str | None = None,
    mime: str | None = None,
    cfg: Mapping[str, Any] | None = None,
) -> AttachmentExtractResult:
    parsed_cfg = AttachmentIngestionConfig.from_cfg(cfg)
    path_obj = Path(path) if path is not None else None
    raw = _load_raw_bytes(path=path_obj, content_bytes=content_bytes)
    if len(raw) > parsed_cfg.max_file_bytes:
        raise ValueError(
            f"attachment exceeds max_file_bytes={parsed_cfg.max_file_bytes}: got={len(raw)}"
        )
    declared = _detect_format(path=path_obj, filename=filename, mime=mime)
    detected = _preflight_attachment(raw, declared_format=declared, cfg=parsed_cfg)
    effective_filename = filename or (path_obj.name if path_obj is not None else None)
    effective_mime = mime
    # Structured parser selection follows verified bytes, not caller-controlled metadata.
    if detected != declared:
        effective_filename = f"attachment.{detected}"
        effective_mime = {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "html": "text/html",
            "image": "image/png",
            "zip": "application/zip",
        }.get(detected, mime)
    kwargs = {
        "path": None,
        "content_bytes": raw,
        "filename": effective_filename,
        "mime": effective_mime,
        "cfg": cfg,
    }
    if not parsed_cfg.sandbox.enabled or detected not in {
        "pdf",
        "docx",
        "html",
        "image",
    }:
        return _extract_attachment_in_process(**kwargs)

    # Native document parsers remain one-shot sandboxed. Images have already
    # passed bounded magic/pixel validation in _preflight_attachment; build their
    # metadata/visual asset locally with OCR disabled, then send bytes only to the
    # persistent resource-bounded OCR worker. This avoids reloading Python/image
    # dependencies in a one-shot parser for every image and makes prewarm real.
    if detected == "image" and str(parsed_cfg.ocr.enabled).strip().lower() != "false":
        base_cfg = _config_with_ocr_disabled(cfg)
        base = _extract_attachment_in_process(
            path=None,
            content_bytes=raw,
            filename=effective_filename,
            mime=effective_mime,
            cfg=base_cfg,
        )
        try:
            ocr_result = _extract_image_ocr_text_from_bytes(
                raw,
                filename=effective_filename,
                mime=effective_mime,
                cfg=parsed_cfg.ocr,
            )
        except Exception:
            ocr_result = _ocr_error_result(parsed_cfg)
        return _merge_image_ocr_result(
            base=base, raw=raw, ocr_result=ocr_result, parsed_cfg=parsed_cfg
        )

    return _extract_attachment_in_sandbox(
        raw=raw,
        filename=effective_filename,
        mime=effective_mime,
        cfg=cfg,
        parsed_cfg=parsed_cfg,
    )


def extract_text_payload(
    *, text: str, cfg: Mapping[str, Any] | None = None
) -> AttachmentExtractResult:
    parsed_cfg = AttachmentIngestionConfig.from_cfg(cfg)
    warnings: List[str] = []
    clipped = _clip_text(
        str(text or ""), max_chars=parsed_cfg.max_extracted_chars, warnings=warnings
    )
    chunks = [
        AttachmentChunk(text=chunk, kind="visible", is_hidden=False)
        for chunk in _chunk_text(
            clipped,
            max_chunk_chars=parsed_cfg.max_chunk_chars,
            chunk_overlap=parsed_cfg.chunk_overlap,
        )
    ]
    full_text = _normalize_text(" ".join(c.text for c in chunks))
    text_empty = len(full_text) == 0
    if text_empty:
        warnings.append("text_empty")
    verdict = "quarantine" if text_empty else "allow"
    return AttachmentExtractResult(
        text=full_text,
        chunks=chunks,
        format="text",
        text_empty=text_empty,
        scan_like=False,
        hidden_text_chars=0,
        warnings=sorted(set(warnings)),
        recommended_verdict=verdict,
        is_image=False,
        image_mime=None,
        image_sha256=None,
        image_bytes_size=0,
        ocr_status="none",
        ocr_provider=None,
        ocr_text_chars=0,
        ocr_spans=[],
    )
